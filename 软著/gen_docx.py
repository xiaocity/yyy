# -*- coding: utf-8 -*-
"""把鉴别材料生成为 Word 文档（.docx）。

分页与取样规则完全复用 gen.py，保证与 PDF 版一致：
源程序取前 30 + 后 30 页、每页 50 行；文档不足 60 页全部提交、每页 30 行。
用显式分页符控制每一页，不依赖 Word 的自动排版——否则行高一变页数就飘。
"""
import io, os, sys, struct

import gen  # 复用同一套分页与取样逻辑

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))

# Word 的版心比 HTML 那套高（页眉在页边距里，不占正文），
# 所以行数一样、行距要单独配，才能同样把页面排满。
# 版心 29.7 - 1.4 - 1.2 = 27.1cm = 768pt，留 16pt 给页码行。
PROG_SPACING = 12.4     # 60 行 x 12.2 = 732pt + 页码 13pt，留 23pt 余量
DOC_SPACING  = 16.9     # 44 行 x 16.6 = 730pt；带图页 30 行留 257pt 给图


def set_font(run, name, size_pt, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 中文字体要单独指定 eastAsia，否则 Word 会回落到默认宋体以外的字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def new_doc(kind=u''):
    """kind 写进页眉尾部，标明这份是哪种鉴别材料（参照件的页眉即为
       「软件全称 版本号 软件说明书」）。源程序件传空串。"""
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21), Cm(29.7)
    s.top_margin, s.bottom_margin = Cm(1.4), Cm(1.2)
    s.left_margin, s.right_margin = Cm(1.2), Cm(1.2)
    s.header_distance, s.footer_distance = Cm(0.7), Cm(0.7)
    # 页眉：软件全称 + 版本号，每一页自动重复
    hp = s.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.tab_stops.add_tab_stop(Cm(18.6), WD_TAB_ALIGNMENT.RIGHT)
    set_font(hp.add_run(gen.NAME), u'宋体', 10.5, bold=True)
    hp.add_run('\t')
    set_font(hp.add_run(gen.VER + (u'　' + kind if kind else u'')), u'宋体', 9.5)

    # 页码在页脚居中。页脚落在下边距里（footer_distance 0.7cm + 行高约 0.46cm
    # < bottom_margin 1.2cm），不占版心，所以行距和每页行数都不用重调。
    fp = s.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(fp.add_run(u'第 '), u'宋体', 9.5)
    add_page_field(fp)                       # Word 的 PAGE 域，自动就是 1..N
    set_font(fp.add_run(u' 页'), u'宋体', 9.5)
    return doc


def add_page_field(paragraph):
    """插入 Word 的 PAGE 域。手写页码要靠自己数，域是 Word 自己算的，
       页数一变就跟着变，不会对不上。"""
    r = paragraph.add_run()
    set_font(r, u'宋体', 9.5)
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = 'PAGE'
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    r._r.append(b); r._r.append(i); r._r.append(e)
    return r


def add_line(doc, text, font, size, spacing_pt, align=None, color=None, bold=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before, pf.space_after = Pt(0), Pt(0)
    pf.line_spacing = Pt(spacing_pt)
    if align is not None:
        p.alignment = align
    set_font(p.add_run(text if text else u''), font, size, bold=bold, color=color)
    return p


def _last_content_is_table(doc):
    """正文最后一个「内容」元素是不是表格。

    不能直接看 body[-1]：w:sectPr（节属性）恒为正文的最后一个子元素，
    body[-1] 永远不等于 w:tbl，判断会一直不成立——分页符照旧插到表格
    前面，图被挤到下一页。必须跳过 sectPr 往前找。
    """
    for child in reversed(doc.element.body):
        if child.tag == qn('w:sectPr'):
            continue
        return child.tag == qn('w:tbl')
    return False


def page_break(doc, last):
    """分页符必须挂在已有段落上：单独开一个空段落来放它的话，
       那个段落自带默认段距，每页会被顶高约 22pt，把内容挤到下一页
       ——实测 60 页变成过 119 页。"""
    if last:
        return
    if _last_content_is_table(doc):
        # 页尾是表格（附录的三联排图）时不能用 doc.paragraphs[-1]：
        # 它不包含表格单元格里的段落，取到的是表格「之前」那一段，
        # 分页符会插在表格前面，把整排图挤到下一页并凭空多出一页
        # ——实测 Word 打开是 22 页，而 HTML 与 PDF 都是 21 页。
        # 这里补一个零行距的空段落来承载分页符，几乎不占高度。
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before, pf.space_after = Pt(0), Pt(0)
        pf.line_spacing = Pt(1)
        p.add_run().add_break(WD_BREAK.PAGE)
        return
    ps = doc.paragraphs
    if ps:
        ps[-1].add_run().add_break(WD_BREAK.PAGE)


FIG_MAX_W_CM = 13.5     # 与 HTML 的 .fig img { max-width:140mm } 对齐
FIG_MAX_H_CM = 6.4      # 与 .fig img { max-height:64mm } 对齐


def png_size(path):
    """读 PNG 的像素宽高（只解 IHDR，不引图像库）。"""
    with io.open(path, 'rb') as f:
        head = f.read(33)
    if head[:8] != b'\x89PNG\r\n\x1a\n':
        return None
    return struct.unpack('>II', head[16:24])


def fit_width_cm(path):
    """按宽、高两个上限等比缩放，返回该用的宽度。

    原来固定按宽 13.5cm 放图。横向裁图（宽高比约 1.86）算出来高 7.3cm 没问题，
    但近正方形的图（设置开关那张是 0.98）会算出 13.8cm 高——HTML 侧有
    max-height:70mm 兜底，docx 侧没有，于是那一页在 Word 里放不下，
    图被挤成单独一页，Word 比 HTML 多出一整页。
    """
    size = png_size(path)
    if not size:
        return FIG_MAX_W_CM
    w, h = size
    return min(FIG_MAX_W_CM, FIG_MAX_H_CM * w / float(h))


def add_figure_row(doc, items, per_row=3):
    """把若干张图排成每行 per_row 列。用无边框表格实现——
       Word 里让图片横向并排只有表格和文本框两条路，表格不依赖浮动定位，
       分页行为可预测，而浮动文本框会在换页时乱跑。
       每张图占两行单元格：上行放图，下行放图注。"""
    width_cm = gen.FIG_ROW_W_MM / 10.0
    chunks = [items[i:i + per_row] for i in range(0, len(items), per_row)]
    t = doc.add_table(rows=2 * len(chunks), cols=per_row)
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 默认表格样式带边框，鉴别材料里会打印成一圈黑线
    t.style = doc.styles['Table Grid']
    set_table_borderless(t)
    col_cm = 18.6 / per_row
    for r, chunk in enumerate(chunks):
        for j, (cap, path) in enumerate(chunk):
            cell = t.cell(r * 2, j)
            cell.width = Cm(col_cm)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6 if r == 0 else 4)
            p.paragraph_format.space_after = Pt(0)
            p.add_run().add_picture(path, width=Cm(width_cm))
            cc = t.cell(r * 2 + 1, j)
            cc.width = Cm(col_cm)
            cp = cc.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(2)
            set_font(cp.add_run(cap), u'宋体', 8.5, color=(0x33, 0x33, 0x33))
    return t


def set_table_borderless(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'none')
        e.set(qn('w:sz'), '0')
        borders.append(e)
    tblPr.append(borders)


def build_source(dst):
    lines = gen.load_source()
    pages = gen.paginate(lines, gen.PROG_LINES_PER_PAGE, merge_tail=False)
    picked, total, cut = gen.pick_30_30(pages)
    doc = new_doc()
    for i, (pageno, rows) in enumerate(picked):
        for r in rows:
            add_line(doc, r, u'Consolas', 8.5, PROG_SPACING)
        page_break(doc, i == len(picked) - 1)
    doc.save(dst)
    return len(picked), total


def build_manual(dst, with_figs=True):
    txt = io.open(os.path.join(HERE, u'软件说明书.txt'), encoding='utf-8').read()
    lines = gen.doc_lines(txt)
    pages, figs = gen.paginate_var(lines, gen.DOC_LINES_PER_PAGE, gen.locate_figs(lines))
    picked, total, cut = gen.pick_30_30(pages)
    doc = new_doc(u'软件说明书')
    for i, (pageno, rows) in enumerate(picked):
        for r in rows:
            add_line(doc, r, u'宋体', 10.5, DOC_SPACING)
        if with_figs and pageno in figs:
            items = [(c, os.path.join(HERE, d, fn)) for c, fn, d in figs[pageno]]
            items = [(c, p) for c, p in items if os.path.exists(p)]
            if len(items) == 1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                p.add_run().add_picture(items[0][1], width=Cm(fit_width_cm(items[0][1])))
                add_line(doc, items[0][0], u'宋体', 9.5, 13,
                         align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x33, 0x33, 0x33))
            elif items:
                add_figure_row(doc, items)
        page_break(doc, i == len(picked) - 1)
    doc.save(dst)
    return len(picked), total


def main():
    figs = '--nofig' not in sys.argv
    n1, t1 = build_source(os.path.join(HERE, u'源程序鉴别材料.docx'))
    print(u'源程序：共 %d 页，提交 %d 页 → 源程序鉴别材料.docx' % (t1, n1))
    n2, t2 = build_manual(os.path.join(HERE, u'文档鉴别材料.docx'), with_figs=figs)
    print(u'文档：共 %d 页，提交 %d 页%s → 文档鉴别材料.docx' % (t2, n2, u'（含界面图）' if figs else u''))


if __name__ == '__main__':
    main()
